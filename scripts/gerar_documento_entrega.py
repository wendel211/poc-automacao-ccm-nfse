"""Gera docs/entrega_desafio_tecnico.docx com margens ABNT, Arial 12, títulos em negrito."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(3)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)


def set_font(run, size=12, bold=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def para(text="", bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold)
    return p


def heading(text, sb=12):
    return para(text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=sb, sa=6)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r)


def table_row_font(cells, size=10, bold=False):
    for cell in cells:
        for run in cell.paragraphs[0].runs:
            run.font.name = "Arial"
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(0, 0, 0)


# ── Cabeçalho ────────────────────────────────────────────────────────────────
for line in ["Wendel Muniz", "wendelmuniz04@gmail.com",
             datetime.date.today().strftime("%d de junho de %Y")]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(line)
    set_font(r)

para(sa=14)

para("Para: Fernanda Batista", sa=2)
para("Assunto: Entrega do Desafio Técnico — POC de Automação CCM + NFS-e", sa=18)

# ── Saudação ─────────────────────────────────────────────────────────────────
para("Prezada Fernanda,", sa=10)
para(
    "Agradeço o retorno. Com base no feedback, evoluí a solução para de fato trazer "
    "resultados, e não apenas levantar os pontos de bloqueio de cada portal. O projeto "
    "é uma automação fiscal em Python que lê a planilha de entrada e, para cada linha, "
    "extrai e valida os dados da nota a partir da própria chave fiscal, enriquece o "
    "cadastro da empresa por APIs públicas e tenta a captura nos portais municipais.",
    sa=10,
)
para(
    "O resultado passou de nenhuma linha com dados para 20 de 25 linhas com dados "
    "completos (cadastro da empresa + documento fiscal decodificado), 5 parciais e "
    "nenhum erro — além de sinalizar 2 divergências reais de CNPJ para auditoria.",
    sa=10,
)

# ── 1. Repositório ────────────────────────────────────────────────────────────
heading("1. Repositório")
para("O código-fonte está disponível publicamente no GitHub:", sa=4)
para("https://github.com/wendel211/poc-automacao-ccm-nfse", sa=10)

# ── 2. Como Executar ──────────────────────────────────────────────────────────
heading("2. Como Executar")
para("Opção 1 — Python local:", bold=True, sa=4)
para(
    'pip install -e ".[dev]"\n'
    "python -m playwright install chromium\n"
    "python -m src.main input/janabril2026_amostra_5x5.xlsx",
    sa=10,
)
para("Opção 2 — Docker (recomendado para reprodutibilidade):", bold=True, sa=4)
para("docker compose run --rm pipeline", sa=10)

# ── 3. O Que o Sistema Faz ───────────────────────────────────────────────────
heading("3. O Que o Sistema Faz")
para("Para cada linha da planilha de entrada, o pipeline:", sa=4)
for item in [
    "Identifica o município, o CNPJ e o código de verificação da nota fiscal.",
    "Seleciona o conector correto por município (Strategy Pattern).",
    "Consulta o CCM — utiliza cache SQLite para evitar consultas duplicadas ao mesmo portal.",
    "Captura screenshot do cadastro municipal como evidência.",
    "Captura screenshot ou acessa o documento da nota fiscal.",
    "Registra o resultado por linha com status, mensagem técnica e caminho dos arquivos.",
    "Ao final, gera planilha de saída com 9 colunas de resultado e relatório HTML com evidências embutidas.",
    "Grava log estruturado JSONL (output/logs/) com uma entrada por linha processada.",
]:
    bullet(item)
para(sa=4)

# ── 4. Municípios ─────────────────────────────────────────────────────────────
heading("4. Municípios Atendidos")
tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for i, h in enumerate(["Município", "Portal", "Estratégia", "Limitação / Observação"]):
    hdr[i].text = h
table_row_font(hdr, bold=True)

for row in [
    ("Belo Horizonte", "servicos.pbh.gov.br", "Playwright (Sydle SPA)", "Screenshot como evidência — Shadow DOM"),
    ("Rio de Janeiro", "Nota Carioca", "Playwright + NFS-e Nacional", "CAPTCHA bloqueia submissão; formulário preenchido capturado"),
    ("Barueri", "ISSNet Online", "Playwright", "Cloudflare 403 — evidência em arquivo .txt"),
    ("Porto Alegre", "NFS-e Nacional", "Chave longa (40+ dígitos)", "DNS failure em todos os portais municipais"),
    ("Nova Lima", "NFS-e Nacional", "Chave longa (40+ dígitos)", "Portal migrado em jan/2026; DNS failure"),
]:
    cells = tbl.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v
    table_row_font(cells)
para(sa=4)

# ── 5. Estratégia de contorno dos bloqueios ──────────────────────────────────
heading("5. Como os Bloqueios dos Portais Foram Contornados")
para(
    "Todos os portais municipais bloqueiam acesso automatizado (CAPTCHA no Rio de "
    "Janeiro, Cloudflare em Barueri, login gov.br no NFS-e Nacional, DNS offline em "
    "Porto Alegre e Nova Lima). Insistir em baixar o PDF no portal não traz resultado. "
    "A virada de estratégia foi extrair os dados diretamente das fontes que de fato "
    "respondem:",
    sa=6,
)
for item in [
    "Decodificação das chaves fiscais: a coluna COD.VERIFICACAO já contém os dados da "
    "nota. As chaves de acesso são padronizadas nacionalmente e auto-contidas — 50 "
    "dígitos para NFS-e Nacional (serviços) e 44 dígitos para NF-e/NFC-e (produtos). "
    "Decodificá-las extrai município, CNPJ do emitente, número, série e competência, "
    "com validação do dígito verificador por módulo 11. Não depende de nenhum portal.",

    "Enriquecimento do cadastro da empresa: como a Inscrição Municipal (CCM) não é "
    "exposta por nenhuma fonte pública federal, o cadastro oficial da empresa é obtido "
    "por uma cadeia de fallback entre APIs públicas (ReceitaWS, CNPJa, BrasilAPI), "
    "retornando razão social, situação cadastral e atividade principal.",

    "Validação cruzada de auditoria: o CNPJ do emitente embutido na chave é comparado "
    "ao fornecedor listado na planilha, sinalizando divergências fiscais.",

    "Tentativas de captura nos portais foram mantidas como evidência (Shadow DOM "
    "piercing em BH, contexto stealth em Barueri, formulário preenchido no RJ).",
]:
    bullet(item)
para(sa=6)

heading("Resultado da Execução Real", sb=6)
para(
    "Executado sobre a planilha de amostra com 25 linhas (5 municípios x 5 empresas):",
    sa=6,
)
tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = "Table Grid"
hdr2 = tbl2.rows[0].cells
for i, h in enumerate(["Status", "Linhas", "Detalhe"]):
    hdr2[i].text = h
table_row_font(hdr2, bold=True)

for row in [
    ("SUCESSO", "20", "Cadastro da empresa obtido + documento fiscal decodificado/capturado"),
    ("PARCIAL", "5",  "Cadastro obtido; código proprietário do portal municipal não decodificável e portal bloqueou captura"),
    ("ERRO",    "0",  "Nenhuma linha sem dados"),
]:
    cells = tbl2.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v
    table_row_font(cells)
para(sa=4)
para(
    "A validação cruzada também sinalizou 2 divergências reais de CNPJ do emitente — "
    "notas cujo emitente difere do fornecedor listado (ex.: nota emitida de Curitiba "
    "para um fornecedor de Barueri).",
    sa=4,
)

# ── 6. Stack ──────────────────────────────────────────────────────────────────
heading("6. Stack Técnica")
for item in [
    "Python 3.11 com Pydantic v2 para validação e normalização de cada linha da planilha.",
    "Decodificador de chaves fiscais próprio (NFS-e Nacional 50 dígitos e NF-e/NFC-e 44 dígitos) com validação de dígito verificador por módulo 11.",
    "httpx + tenacity para requisições HTTP com retry; cadeia de fallback de enriquecimento de CNPJ (ReceitaWS, CNPJa, BrasilAPI).",
    "Playwright (Chromium headless) com Shadow DOM piercing e contexto stealth para tentativa de captura nos portais.",
    "Rich para tabela de progresso ao vivo no terminal durante a execução.",
    "SQLite para cache de CCM por municipio::cnpj, evitando consultas duplicadas.",
    "Docker / docker-compose para execução reproduzível em qualquer ambiente.",
    "GitHub Actions (CI) para execução automática dos 28 testes a cada push.",
    "openpyxl / pandas para leitura da planilha de entrada e escrita dos resultados.",
    "loguru para logs estruturados; JSONL por execução para auditoria linha a linha.",
]:
    bullet(item)
para(sa=4)

# ── 7. Arquivos Gerados ───────────────────────────────────────────────────────
heading("7. Arquivos Gerados por Execução")
for item in [
    "output/resultado_<timestamp>.xlsx — planilha com 19 colunas de resultado (dados cadastrais, dados da chave decodificada, validações e caminhos de arquivo), coloridas por status.",
    "output/relatorio_<timestamp>.html — relatório HTML com cards de resumo e screenshots embutidos.",
    "output/logs/execution_<timestamp>.jsonl — log estruturado com uma entrada JSON por linha processada.",
    "output/evidencias/<MUNICIPIO>/<CNPJ>/ — screenshots de cadastro e notas fiscais por empresa.",
]:
    bullet(item)
para(sa=4)

# ── 8. Decisões Técnicas ─────────────────────────────────────────────────────
heading("8. Principais Decisões Técnicas")
for item in [
    "Verificação de URLs por HTTP antes de implementar qualquer automação, evitando trabalho em "
    "endpoints inexistentes (DNS failure documentado para Porto Alegre e Nova Lima).",

    "Strategy Pattern nos conectores: o orquestrador não conhece o portal — apenas chama o "
    "conector do município. Permite adicionar novos municípios sem alterar o pipeline.",

    "Detecção de chave NFS-e Nacional via regex \\d{40,}: chaves com 40 ou mais dígitos "
    "numéricos são roteadas ao portal nacional; códigos curtos ao portal municipal.",

    'wait_until="domcontentloaded" no Playwright: portais com CDN e analytics externos nunca '
    "atingem networkidle dentro de 30 segundos, gerando falsos erros de navegação.",

    "Evidência .txt para Barueri: a CSP do Cloudflare bloqueia a captura de screenshot via CDP; "
    "o sistema salva URL, status HTTP e descrição do bloqueio como evidência da tentativa.",
]:
    bullet(item)
para(sa=14)

# ── Encerramento ──────────────────────────────────────────────────────────────
para(
    "Agradeço a oportunidade e fico à disposição para quaisquer dúvidas ou esclarecimentos adicionais.",
    sa=18,
)
para("Atenciosamente,", sa=4)
para("Wendel Muniz", bold=True, sa=2)
para("wendelmuniz04@gmail.com", sa=0)

doc.save("docs/entrega_desafio_tecnico.docx")
print("ok")
